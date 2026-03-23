from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_college_report(structured: dict, output_path: str) -> str:
    """
    Generates a Word document with specific College Report formatting rules.
    """
    doc = Document()
    
    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    
    # Simple Footer
    try:
        footer_para = section.footer.paragraphs[0]
    except IndexError:
        footer_para = section.footer.add_paragraph()
        
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Page numbers are typically managed natively in Word.")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(10)

    # Title
    title = structured.get('title', 'Document Title')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run(title)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(20)
    title_run.bold = True

    # Sections
    for sec in structured.get('sections', []):
        heading = sec.get('heading', '')
        content = sec.get('content', '')

        # Heading
        h_para = doc.add_paragraph()
        h_para.paragraph_format.space_before = Pt(12)
        h_para.paragraph_format.space_after = Pt(6)
        h_run = h_para.add_run(heading)
        h_run.font.name = 'Times New Roman'
        h_run.font.size = Pt(14)
        h_run.bold = True
        h_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Body
        b_para = doc.add_paragraph(content)
        b_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        b_para.paragraph_format.line_spacing = 1.5
        b_para.paragraph_format.space_after = Pt(6)
        
        # Apply font to body paragraph
        for run in b_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.save(output_path)
    return output_path
