from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_order_in_council(structured: dict, output_path: str) -> str:
    """
    Generates a highly formal legal document with strict instructions.
    """
    doc = Document()
    
    # Page Setup (A4 or Letter, 1.5 left margin, 1 in else)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # Header logic is somewhat tricky natively, but 
    # title can act as the top title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("ORDER IN COUNCIL")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(12)
    
    # Reference / Subtitle
    ref_title = structured.get('title', 'Reference Document')
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(ref_title)
    ref_run.font.name = 'Times New Roman'
    ref_run.font.size = Pt(14)
    ref_run.bold = True
    ref_para.paragraph_format.space_after = Pt(24)

    # Content
    section_number = 1
    for sec in structured.get('sections', []):
        heading = sec.get('heading', '')
        content = sec.get('content', '')

        # Numbered Heading
        h_para = doc.add_paragraph()
        h_run = h_para.add_run(f"{section_number}. {heading}")
        h_run.font.name = 'Times New Roman'
        h_run.font.size = Pt(14)
        h_run.bold = True
        h_para.paragraph_format.space_after = Pt(12)
        h_para.paragraph_format.space_before = Pt(12)
        
        # Content body - Justified, no indent, 1.5 line spacing, 14pt
        b_para = doc.add_paragraph(content)
        b_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        b_para.paragraph_format.line_spacing = 1.5
        # Full line break after paragraph
        b_para.paragraph_format.space_after = Pt(21)

        for run in b_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            
        section_number += 1

    # Signature Block
    doc.add_paragraph() # extra spacing
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.paragraph_format.space_before = Pt(40)
    sig_run = sig_para.add_run("_____________________________\nSignature")
    sig_run.font.name = 'Times New Roman'
    sig_run.font.size = Pt(14)

    # Footer (Page number usually requires complex Oxml manipulation in python-docx,
    # so we add a simple text instead)
    try:
        footer_para = section.footer.paragraphs[0]
    except IndexError:
        footer_para = section.footer.add_paragraph()
        
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_para.add_run("Page numbers generated natively during print")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(10)

    doc.save(output_path)
    return output_path
