from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_proof_of_concept(structured: dict, output_path: str) -> str:
    """
    Generates a Word document with professional business formatting.
    """
    doc = Document()
    
    # Title
    title = structured.get('title', 'Proof of Concept')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(24)

    for sec in structured.get('sections', []):
        heading = sec.get('heading', '')
        content = sec.get('content', '')

        # Heading
        h_para = doc.add_paragraph()
        h_run = h_para.add_run(heading.upper())
        h_run.font.name = 'Calibri Light'
        h_run.font.size = Pt(16)
        h_run.bold = True
        h_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        h_para.paragraph_format.space_before = Pt(18)
        h_para.paragraph_format.space_after = Pt(6)
        
        # Content
        b_para = doc.add_paragraph(content)
        b_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        b_para.paragraph_format.space_after = Pt(12)
        for run in b_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    doc.save(output_path)
    return output_path
