from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_technical_readme(structured: dict, output_path: str) -> str:
    """
    Generates a Word document formatted as a Technical README.
    """
    doc = Document()
    
    # Title
    title = structured.get('title', 'Technical README')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.name = 'Consolas'
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    doc.add_paragraph() # Spacing

    for sec in structured.get('sections', []):
        heading = sec.get('heading', '')
        content = sec.get('content', '')

        # Heading
        h_para = doc.add_paragraph()
        h_run = h_para.add_run(f"## {heading}")
        h_run.font.name = 'Consolas'
        h_run.font.size = Pt(16)
        h_run.bold = True
        
        # Content
        b_para = doc.add_paragraph(content)
        b_para.paragraph_format.space_after = Pt(12)
        for run in b_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.save(output_path)
    return output_path
