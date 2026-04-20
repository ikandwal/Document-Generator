from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import datetime

def generate_research_paper(structured: dict, output_path: str, subtype: str = None) -> str:
    """
    Generates a Word document with academic Research Paper formatting.
    """
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    title = structured.get('title', 'Research Paper')

    # Subtype or header
    if subtype:
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run(subtype)
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(10)
        header_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(36)
    title_para.paragraph_format.space_after = Pt(24)
    title_run = title_para.add_run(title)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # Sections
    for sec in structured.get('sections', []):
        heading = sec.get('heading', '')
        content = sec.get('content', '')

        # Filter out redundant components
        if heading.lower().strip() == 'title':
            continue

        h_para = doc.add_paragraph()
        h_para.paragraph_format.space_before = Pt(18)
        h_para.paragraph_format.space_after = Pt(6)
        
        # Center the Abstract heading, left-align others
        if heading.lower().strip() == 'abstract':
            h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        h_run = h_para.add_run(heading)
        h_run.font.name = 'Times New Roman'
        h_run.font.size = Pt(12)
        h_run.bold = True
        h_run.font.color.rgb = RGBColor(0, 0, 0)
        
        b_para = doc.add_paragraph(content)
        b_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        b_para.paragraph_format.line_spacing = 1.5
        b_para.paragraph_format.space_after = Pt(12)
        
        for run in b_para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    doc.save(output_path)
    return output_path
