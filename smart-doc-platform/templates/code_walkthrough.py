from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_code_walkthrough(structured: dict, output_path: str) -> str:
    """
    Generates a Word document with step-by-step code walkthrough formatting.
    """
    doc = Document()
    
    # Title
    title = structured.get('title', 'Code Walkthrough')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run(title)
    title_run.font.name = 'Helvetica'
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_para.paragraph_format.space_after = Pt(20)

    step_number = 1
    for sec in structured.get('sections', []):
        heading = sec.get('heading', '')
        content = sec.get('content', '')

        # Step Heading
        h_para = doc.add_paragraph()
        h_run = h_para.add_run(f"Step {step_number}: {heading}")
        h_run.font.name = 'Helvetica'
        h_run.font.size = Pt(14)
        h_run.bold = True
        h_run.font.color.rgb = RGBColor(0x00, 0x4B, 0x87)
        h_para.paragraph_format.space_after = Pt(6)
        
        # Content
        b_para = doc.add_paragraph(content)
        b_para.paragraph_format.space_after = Pt(18)
        b_para.paragraph_format.left_indent = Inches(0.25)
        for run in b_para.runs:
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
            
        step_number += 1

    doc.save(output_path)
    return output_path
