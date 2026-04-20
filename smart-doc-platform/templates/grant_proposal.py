from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import datetime

def generate_grant_proposal(structured: dict, output_path: str, subtype: str = None) -> str:
    """
    Generates a Word document for a Grant Proposal.
    """
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    title = structured.get('title', 'Grant Proposal')
    
    # Subtype or header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run(subtype if subtype else "Grant Proposal")
    header_run.font.name = 'Arial'
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Title Page content
    doc.add_paragraph()
    doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(24)
    title_run = title_para.add_run(title)
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(33, 53, 71)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(48)
    date_run = date_para.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(12)

    # Sections
    for sec in structured.get('sections', []):
        heading = sec.get('heading', '')
        content = sec.get('content', '')

        # Filter out anything that was already rendered
        if heading.lower().strip() == 'title page' or heading.lower().strip() == 'title':
            continue

        h_para = doc.add_paragraph()
        h_para.paragraph_format.space_before = Pt(18)
        h_para.paragraph_format.space_after = Pt(6)
        h_run = h_para.add_run(heading.upper())
        h_run.font.name = 'Arial'
        h_run.font.size = Pt(14)
        h_run.bold = True
        h_run.font.color.rgb = RGBColor(41, 128, 185) # Blue tint for persuasive clear look
        
        b_para = doc.add_paragraph(content)
        b_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        b_para.paragraph_format.line_spacing = 1.15
        b_para.paragraph_format.space_after = Pt(12)
        
        for run in b_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    doc.save(output_path)
    return output_path
